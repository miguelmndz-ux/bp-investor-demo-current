"""Append the post-audit strategy discussion to the existing audit docx."""
import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DOC_PATH = Path(__file__).parent / "BuildParty-Frontend-Audit-2026-04-23.docx"

PRIMARY = RGBColor(0x9C, 0x3F, 0x00)
ACCENT = RGBColor(0xC2, 0x4E, 0x00)
MUTED = RGBColor(0x7F, 0x51, 0x2E)
P0 = RGBColor(0xB0, 0x25, 0x00)
P1 = RGBColor(0xC2, 0x4E, 0x00)
P2 = RGBColor(0x7A, 0x54, 0x00)
GREEN = RGBColor(0x2E, 0x7D, 0x32)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color is not None:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_mixed_para(doc, parts, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for text, fmt in parts:
        run = p.add_run(text)
        run.bold = fmt.get('bold', False)
        run.italic = fmt.get('italic', False)
        if 'color' in fmt:
            run.font.color.rgb = fmt['color']
        if 'size' in fmt:
            run.font.size = Pt(fmt['size'])
        if fmt.get('mono'):
            run.font.name = 'Consolas'
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


def add_labeled_bullet(doc, label, body, label_color=ACCENT):
    p = doc.add_paragraph(style='List Bullet')
    lr = p.add_run(label + " — ")
    lr.bold = True
    lr.font.color.rgb = label_color
    p.add_run(body)
    return p


# ============================================================
# OPEN & APPEND
# ============================================================

doc = Document(str(DOC_PATH))

# Page break to separate strategy from audit
doc.add_page_break()

# ------------------------------------------------------------
# PART 2 — STRATEGY
# ------------------------------------------------------------
title = doc.add_paragraph()
tr = title.add_run("Part 2 — Post-Audit Strategy")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = PRIMARY

subtitle = doc.add_paragraph()
sr = subtitle.add_run("Production path, Figma-first workflow, and the 3-week investor-demo plan")
sr.italic = True
sr.font.size = Pt(12)
sr.font.color.rgb = MUTED

doc.add_paragraph()

# ------------------------------------------------------------
# Q1 — Do we start over?
# ------------------------------------------------------------
add_heading(doc, "Q1. Do we need to rebuild from scratch?", level=1, color=PRIMARY)
add_mixed_para(doc, [
    ("Short answer: ", {'bold': True}),
    ("No. Starting over would throw away the architecture (AppShell, useIsMobile, shared primitives, fixture layer), test coverage for every component, and the hard-won gotchas documented in CLAUDE.md, all to solve problems that are mostly find-and-replace mechanical. The Tailwind tokens are already defined; they're just not consistently used, which is a migration, not a rewrite.", {}),
])

add_mixed_para(doc, [
    ("Honest tradeoff: ", {'bold': True}),
    ("The BARE_ROUTES signup funnel (signup, profile-details, slot-selection, session-confirmation) is the one area where a rewrite might beat a retrofit — those pages have hardcoded px widths and bypass the responsive shell entirely. Keep the investor demo intact, harden it page-by-page, and consider rewriting those 4-5 signup pages mobile-first when you're ready to ship them to real users.", {}),
])

# ------------------------------------------------------------
# Q2 — Changing brand colors
# ------------------------------------------------------------
add_heading(doc, "Q2. What if I want to redefine the visual design / change brand colors?", level=1, color=PRIMARY)
doc.add_paragraph(
    "The hardening work the audit recommends (extracting liquid-glass, distilling gradient-text, "
    "migrating 148 hex values to tokens) is the same work that enables a brand change. Do it in this "
    "order:"
)
for i, step in enumerate([
    "Consolidate duplicated patterns into primitives/classes (<GradientText>, .liquid-glass, extended PillButton).",
    "Migrate inline hexes to reference Tailwind tokens.",
    "At that point, a full rebrand is effectively a tailwind.config.ts edit plus a few CSS custom properties in globals.css.",
], 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)

add_mixed_para(doc, [
    ("Tradeoff: ", {'bold': True}),
    ("If you swap brand colors first (before distilling), you'll grind through ~148 inline hex sites × 9 orange variants and still end up with drift. If you invest a week in the tokenization pass, the brand becomes a config knob. Treat \"tokenize the codebase\" and \"redesign the brand\" as one combined workstream, with tokenization as the enabling step.", {}),
])

# ------------------------------------------------------------
# Q3 — Figma-first workflow
# ------------------------------------------------------------
add_heading(doc, "Q3. What if I design the system in Figma first and import tokens to code?", level=1, color=PRIMARY)
doc.add_paragraph(
    "Great approach, and it doesn't change the code work you eventually have to do — it just ensures "
    "when you do it, you're migrating to the right tokens instead of the current ones. The workflow:"
)
for step in [
    "Design tokens + components in Figma (color, type, spacing, radius, shadow, components, variants).",
    "Export via Figma Variables REST API or Tokens Studio (plugin with GitHub sync as JSON).",
    "Transform through Style Dictionary (or custom script) to generate tailwind.config.ts + CSS custom properties.",
    "Then run the /distill + token-migration pass in code against the generated config.",
]:
    p = doc.add_paragraph(style='List Number')
    p.add_run(step)

add_mixed_para(doc, [
    ("Key insight: ", {'bold': True}),
    ("Figma-first buys you faster visual iteration and a Rob-approvable source of truth. It does not shortcut the ~148 inline-hex replacements in code — it just means you only do that migration once, against tokens you actually believe in, rather than twice.", {}),
])

# ------------------------------------------------------------
# Recommended approach (pre-deadline)
# ------------------------------------------------------------
add_heading(doc, "Recommended approach (before the 3-week constraint)", level=1, color=PRIMARY)
doc.add_paragraph(
    "This was the ideal plan before factoring in the investor conference. It's preserved here because "
    "it's still the right post-conference plan."
)

add_heading(doc, "Phase 0 — Set the new code rules (this week)", level=2, color=ACCENT)
doc.add_paragraph(
    "Write a short 'new code rules' addendum to CLAUDE.md: no inline hex colors, no new raw <img>, "
    "every icon button gets an aria-label, use PillButton/OutlineButton/PrimaryButton or extend them. "
    "Stops the debt from growing while the rest of the work is in flight."
)

add_heading(doc, "Phase 1 — Design system in Figma (you, ~1-2 weeks)", level=2, color=ACCENT)
doc.add_paragraph(
    "Build tokens, core components, and variants in Figma. Get Rob to sign off on the palette and "
    "primitives before any token gets imported to code. This is the only phase where visual decisions "
    "get made — everything downstream is execution."
)

add_heading(doc, "Phase 2 — Code hardening that doesn't depend on brand (parallel to Phase 1)", level=2, color=ACCENT)
doc.add_paragraph("Safe to do against current colors, because it's structural, not visual:")
add_labeled_bullet(doc, "/distill",
    "extract liquid-glass, <GradientText>, adopt PillButton everywhere. Duplication collapses from 15× to 1×.")
add_labeled_bullet(doc, "/harden",
    "modal dialog semantics, aria-labels on icon buttons, focus-visible rings, form label associations.")
add_labeled_bullet(doc, "/optimize",
    "migrate 13 <img> → next/image, replace render-blocking CDN fonts with next/font/local, debounce the carousel resize listener.")
doc.add_paragraph(
    "After this, the codebase is production-grade mechanically — it just wears the old skin."
)

add_heading(doc, "Phase 3 — Import Figma tokens → regenerate config", level=2, color=ACCENT)
doc.add_paragraph(
    "Use Tokens Studio (GitHub-sync) + Style Dictionary to generate tailwind.config.ts and CSS custom "
    "properties from Figma. Because Phase 2 already consolidated patterns into primitives, swapping "
    "the palette is ~a single PR that touches config files and the extracted primitives. Everything "
    "downstream theme-shifts automatically."
)

add_heading(doc, "Phase 4 — Mobile pass on the signup funnel", level=2, color=ACCENT)
doc.add_paragraph(
    "With the new design system in place, rebuild the four BARE_ROUTES pages mobile-first against the "
    "new tokens. This is the one area where rewrite > retrofit — the hardcoded 46% / 360px widths "
    "aren't worth saving."
)

add_heading(doc, "Phase 5 — /polish + /audit rerun", level=2, color=ACCENT)
doc.add_paragraph("Final pass, re-run the audit, ship the updated demo.")

add_mixed_para(doc, [
    ("Key insight: ", {'bold': True}),
    ("Phases 1 and 2 run in parallel. You design in Figma while the code gets structurally ready to consume whatever palette you land on. Phase 3 is then a merge point, not a migration. No work gets thrown away, and the brand becomes a config file the moment you're ready.", {}),
])

# ------------------------------------------------------------
# 3-WEEK INVESTOR PLAN
# ------------------------------------------------------------
doc.add_page_break()
add_heading(doc, "3-Week Pre-Conference Plan (ACTIVE)", level=1, color=P0)

add_mixed_para(doc, [
    ("Context: ", {'bold': True}),
    ("Rob is showing BuildParty to real investors in 3 weeks and needs (a) all main screens functioning as a faked-data 'functional app' on his laptop, and (b) a basic mobile-responsive version so he can pull it up on his phone in the wild.", {}),
])

# Mobile scope box
add_heading(doc, "Mobile scope: \"phone-demo-grade,\" not \"phone-user-grade\"", level=2, color=ACCENT)
doc.add_paragraph("Rob opening BuildParty on his phone over coffee ≠ a real user signing up on mobile. Target:")
add_bullet(doc, "No horizontal scroll on any demoed page.")
add_bullet(doc, "Tap targets that don't frustrate.")
add_bullet(doc, "Narrative still reads on a 390px screen.")
add_bullet(doc, "Explicitly OK: some layouts look cramped, some features are hidden/deferred on mobile.")

# Mobile priority pages
add_heading(doc, "Priority pages for mobile (in order)", level=2, color=ACCENT)
mobile_pages = [
    ("/apex", "dashboard + scan overlay. Money shot."),
    ("/discover", "marketplace story."),
    ("/apex/community/[slug]/owner", "core Apex → community story."),
    ("/agents", "intro screen."),
    ("Navigation", "TopNav + BottomNav already exist; verify they don't collide."),
]
for route, desc in mobile_pages:
    p = doc.add_paragraph(style='List Number')
    rr = p.add_run(route + " — ")
    rr.bold = True
    rr.font.name = "Consolas"
    p.add_run(desc)

add_mixed_para(doc, [
    ("Deferred on mobile (on purpose): ", {'bold': True}),
    ("signup funnel (Rob won't demo signup on phone), PreParty live session (too much layout rework — either cut from mobile or show a 'best on desktop' hint), any sub-community pages beyond the primary one.", {}),
])

# The actual 3 weeks
add_heading(doc, "The 3 weeks", level=2, color=ACCENT)

add_heading(doc, "Week 1 — Mobile pass on the top 4 pages", level=3)
doc.add_paragraph(
    "The Apex scan overlay is the hardest — it's a desktop-sized choreography. Decide early: either "
    "rebuild it mobile-first, or gate it behind useIsMobile and show a simplified 'Apex is running…' "
    "card on phone. The second option is 10× cheaper and probably fine for a demo. Same decision for "
    "PreParty."
)

add_heading(doc, "Week 2 — Finish mobile + dress rehearsal (phone AND laptop)", level=3)
doc.add_paragraph(
    "Rob walks through the whole narrative on his phone, then on his laptop. Note every stutter, "
    "overflow, dead link, weird transition. Fix only what breaks the narrative. Wire the four "
    "href=\"#\" links (P1-3) to 'coming soon' stubs — investors will click things, and dead anchors "
    "are the one a11y finding that's also a demo-killer."
)

add_heading(doc, "Week 3 — Polish + final dress rehearsal", level=3)
doc.add_paragraph(
    "Narrow sweep on animations and state transitions. Fix the OutreachDraftModal keyboard trap "
    "(P0-1) only because Rob might tab by accident mid-demo and break the illusion. Final "
    "run-through on both devices."
)

# What's cut
add_heading(doc, "What's cut (and will be picked up after the conference)", level=2, color=P0)
cuts = [
    ("No rebrand.", "A half-done visual redesign will look worse than the current consistent one. Lock the current palette until after the conference."),
    ("No full token migration.", "The 148 inline hexes are invisible to investors."),
    ("No a11y sweep.", "Investors watching Rob screen-share don't use screen readers. (P0 modal exception above.)"),
    ("No responsive signup funnel.", "BARE_ROUTES pages stay desktop-only for the demo; picked up post-conference."),
    ("No full 'distill' refactor.", "Painful, but 3 weeks can't absorb it. Apply Phase 0 rules to new code only."),
    ("No Figma design system build.", "It's a 2-week effort minimum and returns zero value for this conference."),
]
for label, body in cuts:
    add_labeled_bullet(doc, label.rstrip('.'), body, label_color=P0)

# Taxonomy update
add_heading(doc, "Taxonomy update (applied 2026-04-23)", level=2, color=GREEN)
doc.add_paragraph("Three screens previously marked 'Not built' are confirmed built and have been updated in CLAUDE.md:")
tax = [
    ("1.2.1 Decode Page", "/apex/community/[slug]/decode"),
    ("5.0.2 You're in! Overlay", "/discover (PreviewPanel overlay)"),
    ("5.0.3 Rob clicks \"Join the PreParty\"", "/discover (PreviewPanel CTA)"),
]
for name, route in tax:
    p = doc.add_paragraph(style='List Bullet')
    nr = p.add_run(name + " — ")
    nr.bold = True
    rr = p.add_run(route)
    rr.font.name = "Consolas"
    rr.font.size = Pt(10)
doc.add_paragraph(
    "Don't forget to propagate the same status change to the FigJam board (file key lHeDvpwvVswnuSVF0tAcao) "
    "and Figma design file (file key Fl5XvddT3QsN2VRokrFuqs) per the CLAUDE.md three-artifact sync rule."
)

# Bottom line
add_heading(doc, "The trade", level=2, color=PRIMARY)
doc.add_paragraph(
    "You're trading future-debt for present-narrative. That's the right trade for a 3-week demo "
    "window. Don't let \"doing it right\" become the enemy of \"Rob closes the round.\" Re-run "
    "/audit after the conference with investor feedback in hand — that'll sharpen brand direction "
    "for the Figma design system work in Phase 1."
)

# Save
doc.save(str(DOC_PATH))
print(f"Appended strategy to: {DOC_PATH}")
print(f"New size: {DOC_PATH.stat().st_size:,} bytes")
