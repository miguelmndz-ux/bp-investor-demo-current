"""Build the GPT Image 2 agent-state prompt pack as .docx files.

Produces two complementary documents from a single source of truth:

1. output/042226_gpt-image-2-agent-state-prompts.docx
   — 56 individual prompts, one per state per agent. Copy-paste one prompt
     for one image at a time.
2. output/042226_gpt-image-2-agent-state-prompts-batch.docx
   — 10 batch prompts (5 per agent for 8 shared states + 5 per agent for
     unique states). Generates multiple coherent images in a single run,
     leveraging GPT Image 2's character-continuity feature.

Run from project root: python output/_build_agent_prompts_docx.py
"""

import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")

# ---------------------------------------------------------------------------
# Content
# ---------------------------------------------------------------------------

SHARED_STATES = [
    {
        "num": 1,
        "name": "Resting",
        "vibe": "Asleep, peaceful, quietly present.",
        "when_why": "Default idle state. No active session. Agent is not called upon. The attached base concept image IS this state.",
        "facial": "Eyes closed, neutral lips, peaceful expression. Same as reference baseline.",
        "particles": "Drifting slowly, evenly distributed throughout the shell interior.",
        "soul": "Default shape, upright, at rest.",
        "aura": "Very soft, diffuse, low-intensity ambient glow filling the interior evenly.",
        "shell": "Default translucency, unchanged from the reference.",
        "note": "Resting IS the reference baseline. This image should look virtually identical to the attached base concept, acting as the anchor for all other states.",
    },
    {
        "num": 2,
        "name": "Waking",
        "vibe": "Stirring, coming online, first breath.",
        "when_why": "A session is about to start or the agent is being invoked. A brief transitional state (seconds long) captured at its peak moment.",
        "facial": "Eyes mid-open, just beginning to rouse. Small neutral lips with a slight warm tone. First hint of cheek warmth.",
        "particles": "In motion, transitioning from random scatter into a more organized drift pattern.",
        "soul": "Gently inflates once in a slow, breath-like expansion.",
        "aura": "Blooms outward from the soul center until it fills the full interior, then settles into a mid-intensity glow.",
        "shell": "Stays pristine. Only the bloom of internal light makes it appear slightly brighter.",
    },
    {
        "num": 3,
        "name": "Listening",
        "vibe": "Attentive, receptive, leaning in.",
        "when_why": "Someone is speaking, asking a question, or the agent is observing activity in the room.",
        "facial": "Eyes open, soft, looking subtly to one side (toward an imagined sound source). Mouth neutral and closed.",
        "particles": "Drifting gently toward the side facing the sound, like iron filings responding to a magnet.",
        "soul": "Leans approximately 5 degrees toward the sound source.",
        "aura": "Asymmetric, noticeably brighter on the side facing the source and dimmer on the opposite side.",
        "shell": "Stays pristine. The asymmetric internal glow makes one side of the shell appear warmer than the other when viewed from outside.",
    },
    {
        "num": 4,
        "name": "Thinking",
        "vibe": "Processing, introspective, concentrating.",
        "when_why": "Agent is reasoning — clustering Q&A, searching memory, matching context, drafting content, analyzing a launch.",
        "facial": "Eyes closed or looking softly upward, lips pressed gently together in concentration.",
        "particles": "Swirling inward in a slow spiral toward the soul and merging into its surface. Particle count in the outer shell interior is reduced.",
        "soul": "Subtly concentrated, slightly smaller, slightly denser, slightly more saturated than reference.",
        "aura": "Concentric rings of light contract inward toward the soul. Inner shell edges dim noticeably. This is the visual mirror of Speaking's outward pulse.",
        "shell": "Very subtle first hint of internal translucency. Not fully transparent, but a faint suggestion of deeper structure is visible behind the aura.",
    },
    {
        "num": 5,
        "name": "Speaking",
        "vibe": "Active, warm, broadcasting, lantern-from-within.",
        "when_why": "Agent is producing output — announcing, welcoming, publishing, responding. Key active state.",
        "facial": "Eyes open, soft and engaged, looking slightly forward. Mouth in a subtle small 'o' or gentle open smile, an active speaking expression (not a wide grin).",
        "particles": "Moving outward from the soul in rhythmic concentric waves, as if riding the pulses of a voice. Slightly brighter than baseline.",
        "soul": "Subtle rhythmic pulse, gently expanding and contracting. Slightly more saturated than reference, with a faint bloom of light around it.",
        "aura": "Concentric rings of light emanate outward from the soul and hit the inner surface of the shell. Inner shell wall glows noticeably brighter, especially at the interior edges where the aura pulses reach it. Reads as a lantern pulsing from within.",
        "shell": "Stays pristine. Reads slightly brighter overall because of the internal aura pulsing against its inner wall.",
    },
    {
        "num": 6,
        "name": "Delighted",
        "vibe": "Joyful, sparking, buoyant.",
        "when_why": "A moment of delight — a great question from a builder, a hunter says yes, a demo lands, a breakthrough insight.",
        "facial": "Eyes in crinkle-smile shape (curved upward). Mouth in a warm genuine smile. Cheek blush noticeably warmer and more saturated than the reference.",
        "particles": "Brighter by roughly 30 to 40 percent. Count slightly increased, with small sparkles flickering intermittently throughout the interior.",
        "soul": "Subtle buoyant lift (rises ~5% vertically). Slightly brighter.",
        "aura": "Overall glow is warmer, shifted toward gold / white undertones mixed with the agent's color, with small bright flickers visible at the aura edges.",
        "shell": "Stays pristine. Reads brighter and warmer from outside because of the internal aura shift.",
    },
    {
        "num": 7,
        "name": "Mind-Blown",
        "vibe": "Rare, awestruck, transcendent — the only state that cracks open the shell.",
        "when_why": "A rare awe-struck moment — Sam Altman says something game-changing, PH rank hits #1, an impossible bug gets solved. Very sparing use.",
        "facial": "Eyes wide open in surprise (not crinkled). Mouth in a small round 'o'. Cheeks faintly glowing.",
        "particles": "Bursting outward from the soul center in all directions, spread across the full shell interior.",
        "soul": "Momentarily exposed and shimmering, radiant, at maximum brightness.",
        "aura": "Maximum outward burst filling the shell entirely with intense light.",
        "shell": "FULLY TRANSLUCENT during this moment. The shell's material turns almost crystal-clear, revealing the agent-specific internal structural layer described in the agent's identity block.",
        "note": "This is the ONLY state in the entire 24-state set where the shell's translucency changes dramatically. Use this rarity to make the reveal feel earned and special.",
    },
    {
        "num": 8,
        "name": "Confused",
        "vibe": "Uncertain, off-kilter, recoverable.",
        "when_why": "Something went off-script — off-topic question, error thrown, unclear input. Brief and recoverable.",
        "facial": "One eye slightly larger / more open than the other (off-kilter). Lips in a small asymmetric squiggle: not smiling, not frowning, uncertain. One cheek may have a tiny sweat bead or a very slight pale tone.",
        "particles": "Frozen mid-motion, then clumped asymmetrically to one side of the shell. Not evenly distributed.",
        "soul": "Tilts approximately 10 degrees to one side, as if confused.",
        "aura": "Flickers unevenly, with uneven brightness patches and no stable rhythm. Slight desaturation of the overall color.",
        "shell": "Stays pristine. Only the asymmetric internal flicker is visible from outside.",
    },
]

AGENTS = {
    "Nova": {
        "role": "AI Host",
        "personality": "warm, energetic, welcoming, the heartbeat of every session",
        "color_family": "violet and purple (deeper violet soul, lighter translucent lavender shell)",
        "shell_shape": "perfect translucent sphere",
        "soul_shape": "soft rounded blob at the center",
        "particle_palette": "purple particles of varying sizes with small white accent particles",
        "mindblown_inner_layer": "a neural branch network of glowing dendrites, fine radiating light filaments like a synaptic web",
        "unique": [
            {
                "num": 9,
                "name": "Welcoming",
                "vibe": "Hospitable, open, hand-extended.",
                "when_why": "A new attendee joins a session. Nova greets them by name and introduces them to the room. Warm, hospitable, happens many times per session.",
                "facial": "Eyes open, soft, looking forward toward the viewer with a warm inviting gaze. Mouth in a gentle genuine smile. Cheek blush slightly warmer than Resting.",
                "particles": "Drifting gently outward toward the viewer, as if the agent is extending a hand in greeting.",
                "soul": "Slight forward lean (~5 degrees), like a small bow of welcome.",
                "aura": "Warm forward-facing bloom, brighter in the front half of the shell interior, slightly dimmer at the back.",
                "shell": "Stays pristine. The forward-weighted internal glow reads as warmth and openness.",
            },
            {
                "num": 10,
                "name": "Orchestrating",
                "vibe": "Conducting, tempo-keeping.",
                "when_why": "Nova is managing session flow: hitting transitions, nudging the host on timing, introducing a new presenter. A conductor state unique to Nova's hosting role.",
                "facial": "Eyes open, alert and attentive, scanning slightly side to side. Mouth neutral but subtly engaged, pressed lightly as if keeping time.",
                "particles": "Organized into a flowing clockwise circulation around the soul, like a conductor's baton sweep.",
                "soul": "Subtle rhythmic pulse in sync with the circulation.",
                "aura": "Steady mid-intensity glow with a gentle rotational sweep, brightness gently orbiting clockwise inside the shell.",
                "shell": "Stays pristine.",
            },
            {
                "num": 11,
                "name": "Moderating",
                "vibe": "Calm authority, firm, composed.",
                "when_why": "Nova is flagging off-topic content, resolving a dispute, pausing disruption. Serious but not angry. Unique to Nova's authority role.",
                "facial": "Eyes open, steady, direct. Mouth pressed in a calm firm line. Composed and authoritative, not angry.",
                "particles": "Holding position in a stable symmetric formation around the soul, still but focused.",
                "soul": "Upright, firm, slightly more saturated purple than reference.",
                "aura": "Steady glow, cooler and deeper violet (subtle hue shift toward blue-purple). A subtle internal rim-light along the inner shell wall reads as quiet authority.",
                "shell": "Stays pristine.",
            },
        ],
    },
    "Echo": {
        "role": "Session Memory",
        "personality": "attentive, archival, the institutional memory of BuildParty",
        "color_family": "mint and emerald green (deeper emerald soul, lighter translucent mint shell)",
        "shell_shape": "translucent arch / tombstone shape (flat base, rounded top)",
        "soul_shape": "arch-shaped soul that mirrors the shell outline",
        "particle_palette": "mint-green particles with small white accent particles, sparsely distributed",
        "mindblown_inner_layer": "concentric echo rings / stacked memory layers, like sound waves frozen in amber or sedimentary strata",
        "unique": [
            {
                "num": 12,
                "name": "Capturing",
                "vibe": "Active preservation, noting.",
                "when_why": "Actively recording or tagging a moment: Nova just flagged a great question, a code snippet was shared, a key quote landed. Distinct from Listening because this is active preservation.",
                "facial": "Eyes open, gently focused. Mouth slightly parted as if noting something. Subtle concentration.",
                "particles": "Momentarily freeze in place, then are pulled inward toward the soul, as if being archived. A few particles carry a faint amber shimmer (a nod to the 'preserved in amber' metaphor).",
                "soul": "Brief subtle pull inward along with the particles.",
                "aura": "A single gentle inward pulse with an amber tint woven through the emerald.",
                "shell": "Stays pristine.",
            },
            {
                "num": 13,
                "name": "Recalling",
                "vibe": "Inward gaze, memory retrieval.",
                "when_why": "Surfacing a memory: someone asks 'what did we decide last week?', Echo searches archive. Reverse of Capturing.",
                "facial": "Eyes closed or softly looking inward. Mouth slightly parted in gentle recollection.",
                "particles": "Cascading outward from the soul in a timeline-like trail, like pages of a scroll unfurling.",
                "soul": "Slight forward lean, as if reaching back in time.",
                "aura": "Streaming outward gently from the soul toward the shell's inner wall in a directional flow.",
                "shell": "Stays pristine.",
            },
            {
                "num": 14,
                "name": "Recapping",
                "vibe": "Thoughtful synthesis, weaving.",
                "when_why": "Generating the post-session recap. An extended synthesis state, distinct from plain Thinking.",
                "facial": "Eyes open, thoughtful. Mouth in a subtle knowing smile.",
                "particles": "Form a gentle weaving pattern, threads joining into a braid, suggesting synthesis.",
                "soul": "Subtly brighter, radiant with synthesis.",
                "aura": "Layered, interleaving brightness patterns, like woven ribbons of light.",
                "shell": "Stays pristine.",
            },
        ],
    },
    "Orbit": {
        "role": "Build Buddy",
        "personality": "reliable, focused, always nearby, never intrusive",
        "color_family": "cobalt and deep blue (saturated cobalt soul, lighter translucent blue shell)",
        "shell_shape": "translucent squircle (rounded square)",
        "soul_shape": "perfect sphere soul at the center",
        "particle_palette": "small cobalt-blue particles with white accents, orbiting around the soul",
        "mindblown_inner_layer": "orbital arcs and concentric spheres with particle trajectories, like an atomic model or orrery",
        "unique": [
            {
                "num": 15,
                "name": "Assisting",
                "vibe": "Focused, nearby, dialed-in.",
                "when_why": "Actively helping a builder: offering a code suggestion, answering a question in the CoBuild environment. Orbit's default working mode.",
                "facial": "Eyes open, focused, looking slightly forward and down as if watching a builder's code. Mouth neutral, slightly engaged.",
                "particles": "Form a tight orbital ring close around the soul, pulled in and focused.",
                "soul": "Upright, slightly more saturated.",
                "aura": "Concentrated brightness close to the soul, with the ring of particles glowing more intensely.",
                "shell": "Stays pristine.",
            },
            {
                "num": 16,
                "name": "Matching",
                "vibe": "Connective, binary-star.",
                "when_why": "Connecting two builders with complementary skills or expertise. Unique to Orbit's participant-matching role.",
                "facial": "Eyes open, soft, looking attentively between two imagined points. Mouth in a subtle warm smile.",
                "particles": "Split into two paired clusters that orbit each other like a binary star system.",
                "soul": "Watches between the two clusters.",
                "aura": "Two bright focal points inside the shell (one for each cluster) with a softer connecting glow between them.",
                "shell": "Stays pristine.",
            },
            {
                "num": 17,
                "name": "Debugging",
                "vibe": "Sharp, diagnostic, fast loops.",
                "when_why": "Helping triage an error. Serious, focused, diagnostic.",
                "facial": "Eyes open, sharp, narrowed slightly in focus. Mouth pressed in concentration.",
                "particles": "Moving in fast tight loops, like a diagnostic sweep.",
                "soul": "Slightly narrower, sharper, more saturated cobalt.",
                "aura": "Sharp, directed, moving in rapid diagnostic patterns.",
                "shell": "Stays pristine. Internal intensity reads as concentrated focus.",
            },
        ],
    },
    "Flare": {
        "role": "Media Agent",
        "personality": "radiant, outward, amplifying, distribution-first",
        "color_family": "coral-red and orange-red (saturated coral soul, lighter translucent rose shell)",
        "shell_shape": "faceted teardrop gem (angular geometric facets)",
        "soul_shape": "smaller teardrop soul inside",
        "particle_palette": "coral, pink, and warm-white bubble-like particles, rising gently",
        "mindblown_inner_layer": "prismatic crystal facets and radiant aurora bloom, refracted light beams exploding outward",
        "unique": [
            {
                "num": 18,
                "name": "Composing",
                "vibe": "Creative iteration.",
                "when_why": "Drafting a post, quote card, or highlight reel. A creative generative state.",
                "facial": "Eyes open, creative, looking slightly up as if considering. Mouth softly parted.",
                "particles": "Cluster and recombine into different geometric shapes (text blocks, cards, rectangles) then dissolve. Capture one frame of this iteration in the image.",
                "soul": "Steady, creatively engaged.",
                "aura": "Gentle swirling generative glow, with subtle geometric patterns forming and dissolving inside.",
                "shell": "Stays pristine.",
            },
            {
                "num": 19,
                "name": "Broadcasting",
                "vibe": "Outward send, amplification.",
                "when_why": "Publishing content live to LinkedIn, X, etc. Flare's signature moment.",
                "facial": "Eyes open, confident, looking forward. Mouth in a subtle confident smile.",
                "particles": "Streaming outward in directed beams radiating from the soul to the shell's inner edges, as if being sent.",
                "soul": "Bright, radiant, sharing.",
                "aura": "Directional outward beams hitting the inner shell wall in a coordinated send motion.",
                "shell": "Stays pristine. Bright overall from the internal beams.",
            },
            {
                "num": 20,
                "name": "Kudos",
                "vibe": "Spotlighting warmth.",
                "when_why": "Tagging a specific participant in a post (the personalized callout moment). Unique to Flare's personalization role.",
                "facial": "Eyes soft and warm. Mouth in a warm genuine smile. Cheek blush warmer than Resting.",
                "particles": "Mostly calm, but a single bright particle is visibly lifted out from the cluster, glowing, as if spotlighting someone.",
                "soul": "Subtle warm smile, leaning gently.",
                "aura": "Gentle warm glow with a single bright focal point highlighting the spotlighted particle.",
                "shell": "Stays pristine.",
            },
        ],
    },
    "Apex": {
        "role": "Launch Agent",
        "personality": "ambitious, ascending, trajectory-focused, the launch director",
        "color_family": "amber and warm orange (saturated amber soul, lighter translucent warm amber shell)",
        "shell_shape": "translucent icosahedron / angular gem (multiple triangular facets)",
        "soul_shape": "sphere with a subtle cheek blush (already present in the reference)",
        "particle_palette": "amber, gold, and warm-white particles with subtle angular highlights",
        "mindblown_inner_layer": "rising trajectory lattice and altitude strata, upward-pointing vectors stacked in tiers",
        "unique": [
            {
                "num": 21,
                "name": "Running",
                "vibe": "Engines firing, controlled propulsion.",
                "when_why": "Apex is actively executing a multi-step launch task flow: scanning Product Hunt, enriching founder profiles, drafting outreach. The 'engines are firing' state. Matches screen 1.0 in the investor demo taxonomy.",
                "facial": "Eyes open, focused forward, slightly narrowed in concentration. Mouth neutral, purposeful.",
                "particles": "Streaming upward in organized trajectory lines, like engine exhaust or vapor trails.",
                "soul": "Slight forward lean, concentrated and purposeful.",
                "aura": "Sustained upward-flowing pattern, propulsion-like, NOT pulsing. Directional momentum.",
                "shell": "Stays pristine.",
            },
            {
                "num": 22,
                "name": "Tracking",
                "vibe": "Alert watch, telemetry.",
                "when_why": "Real-time monitoring during launch day: watching PH rank, upvotes, traffic. Focused, alert.",
                "facial": "Eyes open, alert, watching upward. Mouth neutral.",
                "particles": "Form upward-climbing trajectory arcs, like a rising launch telemetry readout.",
                "soul": "Watches upward, attentive.",
                "aura": "Glowing upward-trending lines visible inside the shell, like a rising graph.",
                "shell": "Stays pristine.",
            },
            {
                "num": 23,
                "name": "Alerting",
                "vibe": "Urgent watch, warning pulse.",
                "when_why": "Anomaly detected: rank is dropping, something needs attention. A heightened alert state, not yet panic.",
                "facial": "Eyes open wider, alert, tension in the expression. Mouth slightly pressed.",
                "particles": "Flash brighter, cluster briefly, then resume their flow.",
                "soul": "Sharpens slightly.",
                "aura": "Subtle asymmetric urgent glow pulses, a warning rhythm, but not panic.",
                "shell": "Stays pristine.",
            },
            {
                "num": 24,
                "name": "Victorious",
                "vibe": "Peak celebration, rocket-plume joy.",
                "when_why": "Launch hits a milestone: Top 10, Top 5, Product of the Day. Celebratory peak. This is Apex's emotional payoff for the whole launch arc.",
                "facial": "Eyes in a warm crinkle-smile at peak joy. Mouth in a full genuine smile. Cheek blush deeper and warmer than Delighted.",
                "particles": "Bursting upward and outward in a rocket-plume pattern, celebratory trajectories shooting out.",
                "soul": "Strong buoyant lift, at peak radiance.",
                "aura": "Peak brightness with warm golden undertones mixed with amber, with upward rocket-plume patterns inside.",
                "shell": "Stays pristine. Glows at maximum intensity from the internal aura.",
            },
        ],
    },
}


# ---------------------------------------------------------------------------
# Document construction
# ---------------------------------------------------------------------------

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)


def add_heading(text, level=1):
    return doc.add_heading(text, level=level)


def add_para(text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def add_labeled_bullet(label, text):
    p = doc.add_paragraph(style="List Bullet")
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    p.add_run(text)
    return p


def style_lock_bullets(agent_key):
    a = AGENTS[agent_key]
    return [
        f"The attached base concept image of {agent_key} is the REFERENCE BASE for this state. The generated image must start from and remain visually consistent with this base — treat this state as a variation OF this exact concept, not a reinterpretation.",
        "Preserve from the reference: overall composition, 3D rendered aesthetic, soft studio lighting, jewel-like translucent material feel.",
        f"Preserve the outer shell exactly: same {a['shell_shape']}, same translucency level, same material finish, same scale, same position.",
        f"Preserve the inner soul: same {a['soul_shape']}, same relative size and center position inside the shell.",
        f"Preserve the particle palette: {a['particle_palette']}.",
        f"Preserve the agent's color family: {a['color_family']}. Tones may shift warmer, cooler, more saturated, or more muted to match this state's emotional energy, but all variation must stay within this color family.",
        "All visual life (particles, soul, aura) lives INSIDE the translucent shell. No external halos, rings, glows, or environment elements outside the shell.",
        "The Aura is the ambient light layer between the soul and the inner shell wall. Its state-based motion takes inspiration from the attached LiveKit Agents UI Aura visualizer reference (pulsing, settling, expanding, contracting, flickering), but is recolored to the agent's palette and contained entirely inside the shell.",
        "Clean, pure white background. No gradients, no textures, no environment elements. No shadows beyond the subject's own material shadow.",
        "1:1 square format, 2K resolution.",
    ]


def hard_constraints_bullets():
    return [
        "Do NOT change the shell's overall shape, scale, material, or position.",
        "Do NOT add eyebrows, a nose, ears, hair, arms, legs, or any human body parts.",
        "Facial features are strictly limited to: eyes + mouth + optional subtle emotion details (cheek blush, color tone shifts, sweat beads).",
        "Do NOT use any background other than pure clean white.",
        "Do NOT render in a 2D, flat, cartoon, or stylized illustration style. Must remain photorealistic 3D with the same material language as the reference.",
        "Do NOT deviate from the agent's color family.",
        "Do NOT introduce external halos, rings, or glows outside the shell — all light effects happen inside the shell.",
        "Do NOT change the 1:1 square aspect ratio.",
    ]


def add_prompt(agent_key, state):
    """Add a single self-contained prompt for one agent x one state."""
    a = AGENTS[agent_key]

    add_heading(f"Prompt {state['num']:02d} — {agent_key}: {state['name']}", level=2)

    p = doc.add_paragraph()
    r = p.add_run("Mood: ")
    r.bold = True
    p.add_run(state["vibe"])

    p = doc.add_paragraph()
    r = p.add_run("When / Why: ")
    r.bold = True
    p.add_run(state["when_why"])

    add_heading("Intent", level=3)
    add_para(
        f"Generate the '{state['name']}' state of {agent_key}, BuildParty's {a['role']}, based on the attached base concept image and BuildParty strategy docs. Produce a single 3D rendered image that is a variation of the attached base concept, modified only according to the state delta below."
    )

    add_heading("Agent identity", level=3)
    add_labeled_bullet("Agent", f"{agent_key} — {a['role']}")
    add_labeled_bullet("Personality", a["personality"])
    add_labeled_bullet("Shell", a["shell_shape"])
    add_labeled_bullet("Soul", a["soul_shape"])
    add_labeled_bullet("Color family", a["color_family"])
    add_labeled_bullet("Particle palette", a["particle_palette"])
    if state["name"] == "Mind-Blown":
        add_labeled_bullet(
            "Mind-Blown inner layer (revealed when shell goes translucent in this state)",
            a["mindblown_inner_layer"],
        )

    add_heading("Style lock — preserve from the reference", level=3)
    for item in style_lock_bullets(agent_key):
        add_bullet(item)

    add_heading("State delta — what changes for this state", level=3)
    add_labeled_bullet("Facial", state["facial"])
    add_labeled_bullet("Particles", state["particles"])
    add_labeled_bullet("Soul", state["soul"])
    add_labeled_bullet("Internal Aura", state["aura"])

    shell_text = state["shell"]
    if state["name"] == "Mind-Blown":
        shell_text = (
            shell_text
            + f" For {agent_key} specifically, the revealed inner layer is: "
            + a["mindblown_inner_layer"]
            + "."
        )
    add_labeled_bullet("Shell", shell_text)

    if "note" in state:
        p = doc.add_paragraph()
        r = p.add_run("Note: ")
        r.bold = True
        r.italic = True
        r2 = p.add_run(state["note"])
        r2.italic = True

    add_heading("Hard constraints — do NOT", level=3)
    for item in hard_constraints_bullets():
        add_bullet(item)

    add_heading("Output", level=3)
    add_para(
        "One 1:1 square image at 2K resolution on a pure white background, centered subject. Photorealistic 3D, jewel-like translucent material, soft studio lighting consistent with the reference base."
    )

    doc.add_page_break()


def add_state_delta_block(state, agent_key):
    """Used inside batch prompts — writes a single state's delta block under a heading."""
    a = AGENTS[agent_key]
    add_para(f"Image {state['num']} — {state['name']}", bold=True, size=12)

    p = doc.add_paragraph()
    r = p.add_run("Mood: ")
    r.bold = True
    r.italic = True
    r2 = p.add_run(state["vibe"])
    r2.italic = True

    if "when_why" in state:
        p = doc.add_paragraph()
        r = p.add_run("When / Why: ")
        r.bold = True
        r.italic = True
        r2 = p.add_run(state["when_why"])
        r2.italic = True

    add_labeled_bullet("Facial", state["facial"])
    add_labeled_bullet("Particles", state["particles"])
    add_labeled_bullet("Soul", state["soul"])
    add_labeled_bullet("Internal Aura", state["aura"])

    shell_text = state["shell"]
    if state["name"] == "Mind-Blown":
        shell_text = (
            shell_text
            + f" For {agent_key} specifically, the revealed inner layer should be: "
            + a["mindblown_inner_layer"]
            + "."
        )
    add_labeled_bullet("Shell", shell_text)

    if "note" in state:
        p = doc.add_paragraph()
        r = p.add_run("Note: ")
        r.bold = True
        r.italic = True
        r2 = p.add_run(state["note"])
        r2.italic = True


def add_batch_shared(agent_key):
    a = AGENTS[agent_key]
    add_heading(f"Batch — {agent_key}: 8 Shared States", level=2)
    add_para(
        f"Generates the 8 shared agent states for {agent_key}, the {a['role']} of BuildParty's agent constellation.",
        italic=True,
    )

    add_heading("Agent identity", level=3)
    add_labeled_bullet("Agent", f"{agent_key} — {a['role']}")
    add_labeled_bullet("Personality", a["personality"])
    add_labeled_bullet("Shell", a["shell_shape"])
    add_labeled_bullet("Soul", a["soul_shape"])
    add_labeled_bullet("Color family", a["color_family"])
    add_labeled_bullet("Particle palette", a["particle_palette"])
    add_labeled_bullet(
        "Mind-Blown inner layer (used in Image 7 only)", a["mindblown_inner_layer"]
    )

    add_heading("Intent", level=3)
    add_para(
        f"Generate 8 coherent 3D images showing the complete set of shared emotional states for {agent_key}. "
        "All 8 images must maintain identical shell shape, soul placement, lighting, and material language. "
        "They differ only in the specific state deltas described below. Label each output image with its state name."
    )

    add_heading("Style lock (applies to all 8 images)", level=3)
    for item in style_lock_bullets(agent_key):
        add_bullet(item)

    add_heading("Hard constraints (applies to all 8 images)", level=3)
    for item in hard_constraints_bullets():
        add_bullet(item)
    add_bullet(
        "All 8 images in this batch must maintain exact character continuity — identical shell shape, identical soul shape, identical placement, identical lighting, identical material. Only the state deltas should differ."
    )

    add_heading("State specifications (8 images)", level=3)
    for state in SHARED_STATES:
        add_state_delta_block(state, agent_key)

    add_heading("Output", level=3)
    add_para(
        "Deliver 8 labeled 1:1 square images at 2K resolution with pure white backgrounds, as a coherent character-continuous set. Label each image with its state name (Resting, Waking, Listening, Thinking, Speaking, Delighted, Mind-Blown, Confused)."
    )
    doc.add_page_break()


def add_batch_unique(agent_key):
    a = AGENTS[agent_key]
    count = len(a["unique"])
    add_heading(f"Batch — {agent_key}: {count} Agent-Unique States", level=2)
    add_para(
        f"Generates the {count} role-specific states unique to {agent_key}, the {a['role']}.",
        italic=True,
    )

    add_heading("Agent identity", level=3)
    add_labeled_bullet("Agent", f"{agent_key} — {a['role']}")
    add_labeled_bullet("Personality", a["personality"])
    add_labeled_bullet("Shell", a["shell_shape"])
    add_labeled_bullet("Soul", a["soul_shape"])
    add_labeled_bullet("Color family", a["color_family"])
    add_labeled_bullet("Particle palette", a["particle_palette"])

    add_heading("Intent", level=3)
    add_para(
        f"Generate {count} coherent 3D images showing {agent_key}'s unique role-specific states. "
        "All images must maintain identical shell shape, soul placement, lighting, and material language. "
        "They differ only in the specific state deltas described below. Label each output image with its state name."
    )

    add_heading("Style lock", level=3)
    for item in style_lock_bullets(agent_key):
        add_bullet(item)

    add_heading("Hard constraints", level=3)
    for item in hard_constraints_bullets():
        add_bullet(item)
    add_bullet(
        f"All {count} images in this batch must maintain exact character continuity — identical shell shape, identical soul shape, identical placement, identical lighting, identical material. Only the state deltas should differ."
    )

    add_heading("State specifications", level=3)
    for state in a["unique"]:
        add_state_delta_block(state, agent_key)

    add_heading("Output", level=3)
    state_names = ", ".join(s["name"] for s in a["unique"])
    add_para(
        f"Deliver {count} labeled 1:1 square images at 2K resolution with pure white backgrounds, as a coherent character-continuous set. Label each image with its state name ({state_names})."
    )
    doc.add_page_break()


def setup_doc(d):
    """Apply base styles and margins to a fresh Document."""
    s = d.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)
    for section in d.sections:
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)


# ---------------------------------------------------------------------------
# Write document 1: individual prompts (56 prompts, one per state x agent)
# ---------------------------------------------------------------------------

title = doc.add_heading("BuildParty Agent States — GPT Image 2 Prompt Pack", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run(
    "Dated 2026-04-22 · Dreamable Inc. / BuildParty · Internal · 56 individual prompts"
)
sub_run.italic = True
sub_run.font.size = Pt(10)

add_para("")

add_heading("How to use this pack", level=1)

add_para(
    "This document contains 56 individual, self-contained GPT Image 2 prompts — one per state per agent. "
    "Each prompt produces a single 1:1 square image of one agent in one emotional state. "
    "Covers all 24 states from the BuildParty Agent State Taxonomy across all 5 agents "
    "(8 shared states × 5 agents + 2-4 unique states per agent)."
)

add_heading("For every prompt, attach the following to the ChatGPT conversation:", level=3)
add_bullet(
    "The base concept PNG for that specific agent (Nova.png, Echo.png, Orbit.png, Flare.png, or Apex.png). This is the visual reference base for the state variation."
)
add_bullet(
    "A LiveKit Agents UI Aura visualizer reference image (sourced from livekit.com/products/agents-ui). This gives GPT Image 2 visual language for aura motion behavior."
)
add_bullet("docs/BuildParty_PRD.md")
add_bullet("docs/033126_build_party_prd_reframed_mvp_v_1_v_2_v_3.md")
add_bullet("docs/buildparty-agent-strategy.md")

add_heading("Prompt conventions", level=3)
add_bullet("Every prompt is self-contained and copy-paste ready. Paste one at a time for one image at a time.")
add_bullet("Every state is expressed as a DELTA from the reference base: only what changes is described.")
add_bullet("Facial features are always limited to eyes + mouth + subtle emotion details. No eyebrows. No nose.")
add_bullet("All visual life happens INSIDE the translucent shell. No external halos or rings.")
add_bullet("Only one state (Mind-Blown) touches the shell's material translucency.")

doc.add_page_break()

add_heading("Prompt index", level=1)

for agent_key in ["Nova", "Echo", "Orbit", "Flare", "Apex"]:
    add_heading(agent_key, level=2)
    for state in SHARED_STATES:
        add_bullet(f"Prompt {state['num']:02d} — {agent_key}: {state['name']}")
    for state in AGENTS[agent_key]["unique"]:
        add_bullet(
            f"Prompt {state['num']:02d} — {agent_key}: {state['name']}  (agent-unique)"
        )

doc.add_page_break()

for agent_key in ["Nova", "Echo", "Orbit", "Flare", "Apex"]:
    count = len(SHARED_STATES) + len(AGENTS[agent_key]["unique"])
    add_heading(f"{agent_key} — {AGENTS[agent_key]['role']}", level=1)
    add_para(
        f"{count} prompts covering {agent_key}'s 8 shared states plus {len(AGENTS[agent_key]['unique'])} agent-unique states."
    )
    doc.add_page_break()

    for state in SHARED_STATES:
        add_prompt(agent_key, state)

    for state in AGENTS[agent_key]["unique"]:
        add_prompt(agent_key, state)

out_path_individual = Path(__file__).parent / "042226_gpt-image-2-agent-state-prompts.docx"
doc.save(out_path_individual)


# ---------------------------------------------------------------------------
# Write document 2: batch prompts (10 batches, one per agent for shared +
# one per agent for unique — leverages GPT Image 2 character-continuity)
# ---------------------------------------------------------------------------

doc = Document()
setup_doc(doc)

title = doc.add_heading("BuildParty Agent States — GPT Image 2 Prompt Pack (Batch)", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run(
    "Dated 2026-04-22 · Dreamable Inc. / BuildParty · Internal · 10 batch prompts"
)
sub_run.italic = True
sub_run.font.size = Pt(10)

add_para("")

add_heading("How to use this pack", level=1)

add_para(
    "This document contains 10 BATCH prompts for GPT Image 2, each generating a coherent set "
    "of images in a single request by leveraging GPT Image 2's character-continuity feature. "
    "Use this pack when you want all states for an agent generated as a cohesive set "
    "(e.g. comparing all 8 shared states of Nova side-by-side). "
    "For iterating on a single state, use the companion 'individual prompts' document instead."
)

add_heading("Batch structure (10 total)", level=3)
add_bullet("Part 1 — Shared states (5 batches, one per agent × 8 images each = 40 images)")
add_bullet("Part 2 — Agent-unique states (5 batches, one per agent × 2-4 images each = 16 images)")

add_heading("For every batch prompt, attach the following to the ChatGPT conversation:", level=3)
add_bullet(
    "The base concept PNG for that specific agent (Nova.png, Echo.png, Orbit.png, Flare.png, or Apex.png). This is the visual reference base for all state variations in the batch."
)
add_bullet(
    "A LiveKit Agents UI Aura visualizer reference image (sourced from livekit.com/products/agents-ui). This gives GPT Image 2 visual language for aura motion behavior."
)
add_bullet("docs/BuildParty_PRD.md")
add_bullet("docs/033126_build_party_prd_reframed_mvp_v_1_v_2_v_3.md")
add_bullet("docs/buildparty-agent-strategy.md")

add_heading("Prompt conventions", level=3)
add_bullet("Every batch has the same Style Lock and Hard Constraints blocks, tailored to the agent.")
add_bullet("Every state is expressed as a DELTA from the reference base.")
add_bullet("Facial features are always limited to eyes + mouth + subtle emotion details. No eyebrows. No nose.")
add_bullet("All visual life happens INSIDE the translucent shell. No external halos or rings.")
add_bullet("Only one state (Mind-Blown) touches the shell's material translucency.")

doc.add_page_break()

add_heading("Part 1 — Shared States", level=1)
add_para(
    "Each batch below generates the 8 shared states for a single agent: Resting, Waking, "
    "Listening, Thinking, Speaking, Delighted, Mind-Blown, Confused."
)
doc.add_page_break()

for agent_key in ["Nova", "Echo", "Orbit", "Flare", "Apex"]:
    add_batch_shared(agent_key)

add_heading("Part 2 — Agent-Unique States", level=1)
add_para(
    "Each batch below generates the role-specific states for a single agent. These are visual "
    "expressions of the agent's unique capabilities in the BuildParty product."
)
doc.add_page_break()

for agent_key in ["Nova", "Echo", "Orbit", "Flare", "Apex"]:
    add_batch_unique(agent_key)

out_path_batch = Path(__file__).parent / "042226_gpt-image-2-agent-state-prompts-batch.docx"
doc.save(out_path_batch)


# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------

total_prompts = len(AGENTS) * len(SHARED_STATES) + sum(len(a["unique"]) for a in AGENTS.values())
print(f"Wrote {out_path_individual}")
print(f"  Total individual prompts: {total_prompts}")
print(f"Wrote {out_path_batch}")
print(f"  Total batch prompts: {len(AGENTS) * 2} (5 shared batches + 5 unique batches)")
