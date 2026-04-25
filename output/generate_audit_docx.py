"""Generate BuildParty frontend audit as a .docx document."""
import sys
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path(__file__).parent / "BuildParty-Frontend-Audit-2026-04-23.docx"

PRIMARY = RGBColor(0x9C, 0x3F, 0x00)
ACCENT = RGBColor(0xC2, 0x4E, 0x00)
MUTED = RGBColor(0x7F, 0x51, 0x2E)
DARK = RGBColor(0x1A, 0x0A, 0x00)
P0 = RGBColor(0xB0, 0x25, 0x00)
P1 = RGBColor(0xC2, 0x4E, 0x00)
P2 = RGBColor(0x7A, 0x54, 0x00)
P3 = RGBColor(0x7F, 0x51, 0x2E)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color is not None:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=None, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_mixed_para(doc, parts, style=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    for text, fmt in parts:
        run = p.add_run(text)
        run.bold = fmt.get('bold', False)
        run.italic = fmt.get('italic', False)
        if 'color' in fmt:
            run.font.color.rgb = fmt['color']
        if 'size' in fmt:
            run.font.size = Pt(fmt['size'])
        if fmt.get('mono'):
            run.font.name = 'Consolas'
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p


def add_finding(doc, tag, tag_color, title, rows):
    p = doc.add_paragraph()
    tag_run = p.add_run(f"[{tag}] ")
    tag_run.bold = True
    tag_run.font.color.rgb = tag_color
    title_run = p.add_run(title)
    title_run.bold = True
    for label, value in rows:
        rp = doc.add_paragraph()
        lr = rp.add_run(f"{label}: ")
        lr.bold = True
        lr.font.color.rgb = MUTED
        vr = rp.add_run(value)
        if label == 'Location':
            vr.font.name = 'Consolas'
            vr.font.size = Pt(10)


# ============================================================
# BUILD DOC
# ============================================================

doc = Document()

# Base style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("BuildParty Frontend Audit")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = PRIMARY

subtitle = doc.add_paragraph()
sr = subtitle.add_run("Investor-demo codebase — production-readiness review")
sr.italic = True
sr.font.size = Pt(12)
sr.font.color.rgb = MUTED

meta = doc.add_paragraph()
meta.add_run("Date: ").bold = True
meta.add_run("2026-04-23    ")
meta.add_run("Branch: ").bold = True
meta.add_run("041826-Miguel    ")
meta.add_run("Scope: ").bold = True
meta.add_run("web/src/**")

doc.add_paragraph()

# ------------------------------------------------------------
# AUDIT HEALTH SCORE
# ------------------------------------------------------------
add_heading(doc, "Audit Health Score", level=1, color=PRIMARY)

table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "#"
hdr[1].text = "Dimension"
hdr[2].text = "Score"
hdr[3].text = "Key Finding"
for cell in hdr:
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
    set_cell_bg(cell, "FFE3D2")

rows_data = [
    ("1", "Accessibility", "1/4",
     "31 findings: 1 P0 (modal lacks dialog semantics), 11 P1 WCAG AA violations, icon buttons without aria-label, inputs without labels, focus indicators removed"),
    ("2", "Performance", "2/4",
     "Render-blocking CDN fonts, 13 raw <img> tags, unthrottled resize listener — offset by correct RAF/timer cleanup"),
    ("3", "Responsive Design", "2/4",
     "BARE_ROUTES signup funnel is desktop-first, PreParty hides chat/channels on mobile with no fallback, hardcoded px widths"),
    ("4", "Theming", "2/4",
     "Token system well-defined but 148 hardcoded hex colors, 9 'orange' variants, liquid-glass and gradient-text duplicated 15× each instead of extracted"),
    ("5", "Anti-Patterns", "3/4",
     "Brand aesthetic leans on AI-slop tells (gradient text, glassmorphism) but deliberate, systematic, and otherwise clean — no emoji, no filler copy, no bento-slop"),
    ("", "Total", "10/20", "Acceptable (significant work needed) — 5/10 on 1–10 scale"),
]
for i, row in enumerate(rows_data, start=1):
    cells = table.rows[i].cells
    for j, value in enumerate(row):
        cells[j].text = value
    if i == 6:
        for cell in cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
            set_cell_bg(cell, "FFEDE4")

doc.add_paragraph()
add_mixed_para(doc, [
    ("On your 1–10 scale: ", {'bold': True}),
    ("5/10", {'bold': True, 'color': P1}),
    (" — strong bones, good architectural decisions, but not a production-ready foundation without hardening.", {}),
])

# ------------------------------------------------------------
# ANTI-PATTERNS VERDICT
# ------------------------------------------------------------
add_heading(doc, "Anti-Patterns Verdict", level=1, color=PRIMARY)

add_mixed_para(doc, [
    ("Does this look AI-generated? ", {'bold': True}),
    ("Partially yes, but intentionally so.", {}),
])
doc.add_paragraph(
    "The orange-gradient-text + premium-glass + VCR-mono-label palette reads as AI-generated-aesthetic "
    "on first glance, but the documentation in CLAUDE.md, the custom concave-corner wedge in "
    "AppShell.tsx:58-82, the pie-chart scan progress dot in SideNav.tsx:102-135, and the deliberate "
    "token system (tailwind.config.ts) show real design thought. This is \"AI aesthetic used as brand "
    "voice\" — a defensible choice for an AI-platform product. It's not random AI slop; it's consistent "
    "AI slop, which is actually a brand."
)
doc.add_paragraph(
    "The real tells to address are discipline failures, not aesthetic ones: the documented patterns "
    "aren't extracted into reusable components/classes, so the aesthetic is copy-pasted 15+ times "
    "instead of referenced once."
)

# ------------------------------------------------------------
# EXECUTIVE SUMMARY
# ------------------------------------------------------------
add_heading(doc, "Executive Summary", level=1, color=PRIMARY)

add_bullet(doc, "Audit Health Score: 10/20 (Acceptable) — 5/10 on your scale")
add_bullet(doc, "~55 total issues found (3 P0, 18 P1, 20 P2, 14 P3)")

add_para(doc, "Top 5 critical issues", bold=True)
top5 = [
    "OutreachDraftModal has no dialog semantics, focus trap, or focus return — the flagship interaction in the Apex flow is unusable for screen-reader and keyboard-only users",
    "Render-blocking CDN fonts (Material Symbols + VCR OSD Mono via <link> in <head>) block FCP/LCP",
    "BARE_ROUTES sign-up funnel is desktop-only — signup/profile-details/slot-selection have hardcoded px widths and no mobile layout; these are the pages real users would actually touch",
    "13 raw <img> tags with no width/height, no lazy-loading, no next/image optimization — will tank CLS and LCP on real devices",
    "Documented design patterns not extracted — liquid-glass, gradient-text, scan-overlay style copied 15× each instead of wrapped in a shared primitive/CSS class; each new screen will deepen this debt",
]
for i, item in enumerate(top5, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(item)

add_mixed_para(doc, [
    ("Recommended next step before building more screens: ", {'bold': True}),
    ("stop adding surface area and run one hardening pass. Every new screen is currently multiplying the accessibility and theming debt linearly.", {}),
])

# ------------------------------------------------------------
# DETAILED FINDINGS — P0
# ------------------------------------------------------------
add_heading(doc, "Detailed Findings by Severity", level=1, color=PRIMARY)

add_heading(doc, "P0 — Blocking", level=2, color=P0)

add_finding(doc, "P0-1", P0, "OutreachDraftModal missing dialog semantics", [
    ("Location", "web/src/components/apex/OutreachDraftModal.tsx:160-170"),
    ("Category", "Accessibility"),
    ("Impact", "Screen readers don't announce it as a modal; no focus trap so Tab escapes into background content; focus doesn't return to the triggering button on close."),
    ("WCAG", "4.1.2 Name, Role, Value (Level A) — fails Level A, not just AA"),
    ("Recommendation", "Add role=\"dialog\", aria-modal=\"true\", aria-labelledby, trap Tab focus, store and restore the opener."),
    ("Suggested command", "/harden"),
])

add_finding(doc, "P0-2", P0, "Render-blocking external stylesheets in root layout", [
    ("Location", "web/src/app/layout.tsx:23-24"),
    ("Category", "Performance"),
    ("Impact", "FCP/LCP penalty on every page load; the whole app waits on fonts.googleapis.com and fonts.cdnfonts.com before first paint."),
    ("Recommendation", "Replace CDN <link> with next/font/local for VCR OSD Mono; drop Material Symbols CDN in favor of Phosphor Icons (already the documented standard — only legacy nav still uses material-symbols)."),
    ("Suggested command", "/optimize"),
])

add_finding(doc, "P0-3", P0, "148 hardcoded hex colors across 7 worst-offender files", [
    ("Location", "web/src/app/what-to-expect/page.tsx (35), ApexScanOverlay.tsx (39), preparty/page.tsx (17), slot-selection/page.tsx (12), session-confirmation/page.tsx (9)"),
    ("Category", "Theming / Anti-Pattern"),
    ("Impact", "Any brand color change requires touching 148 sites; impossible to theme or skin; makes consistency audits impractical."),
    ("Recommendation", "Migrate to design tokens. Start with the 9 orange variants (#9c3f00, #7a2e00, #c24e00, #ff7a2f, #a33800, #9f6c47 vs #9e6b47) — consolidate to the defined Tailwind tokens or add missing variants to the config."),
    ("Suggested command", "/polish (needs a systematic token migration pass)"),
])

# ------------------------------------------------------------
# P1
# ------------------------------------------------------------
add_heading(doc, "P1 — Major (must fix before building more screens)", level=2, color=P1)

p1_findings = [
    ("P1-1", "11 icon-only buttons missing aria-label",
     "TopNav.tsx:48 (Bell), OutreachDraftModal.tsx:168 (close X), PreviewPanel.tsx:25-31 (collapse), slot-selection/page.tsx:211-241 (prev/next month), slot-selection/page.tsx:399-434 (time slots), PreviewPanel.tsx:322-335 (mobile Copy Link / Event Page)",
     "Accessibility",
     "4.1.2 Name, Role, Value (Level A)",
     "Add aria-label to every icon-only <button>; the pattern is already used correctly in TopNav.tsx:42 for mobile search.",
     "/harden"),
    ("P1-2", "Form inputs without associated labels",
     "profile-details/page.tsx:186-240 (5 inputs; styled <label> tags not linked via htmlFor), TopNav.tsx:28-32 (desktop search), signup/page.tsx:272-289 (readOnly email)",
     "Accessibility",
     "1.3.1 Info and Relationships (A), 3.3.2 Labels or Instructions (A)",
     "Add htmlFor/id pairs, or aria-label if the visible label is intentionally omitted.",
     "/harden"),
    ("P1-3", "Dead href=\"#\" links across nav and forms",
     "SideNav.tsx:172 (Calendar), BottomNav.tsx:32 (Calendar), signup/page.tsx:336 (Log in), profile-details/page.tsx:275 (Change email)",
     "Accessibility / UX",
     "—",
     "Route to a stub 'coming soon' page (the documented pattern in CLAUDE.md) or render as disabled <button> with explanation.",
     "/clarify"),
    ("P1-4", "Focus indicators removed via inline outline:'none'",
     "profile-details/page.tsx:186-240, slot-selection/page.tsx:310, slot-selection/page.tsx:399-434, TopNav.tsx:31",
     "Accessibility",
     "2.4.7 Focus Visible (AA)",
     "Replace inline outline:none with focus-visible:ring-2 focus-visible:ring-primary utility class.",
     "/harden"),
    ("P1-5", "13 raw <img> tags without next/image",
     "SideNav.tsx:95, TopNav.tsx:55, ApexScanOverlay.tsx:182,218,270, CommunityCard.tsx:20,36, SessionCard.tsx:28, FounderCard.tsx:13, PhProductRow.tsx, AgentCard.tsx, + more in carousels/cards",
     "Performance",
     "—",
     "Migrate to next/image. Configure remotePatterns in next.config for avatar/Unsplash sources.",
     "/optimize"),
    ("P1-6", "Unthrottled resize listener on carousel",
     "SessionCarousel.tsx:65",
     "Performance",
     "—",
     "Debounce with requestAnimationFrame or extend the existing ResizeObserver (line 66) to cover window-level resize.",
     "/optimize"),
    ("P1-7", "BARE_ROUTES pages are desktop-first",
     "signup/page.tsx, profile-details/page.tsx:24 (hardcoded 46% column width), profile-details/page.tsx:58 (hardcoded 360px card), slot-selection/page.tsx:71, session-confirmation/page.tsx",
     "Responsive",
     "—",
     "Partial mobile support exists via useIsMobile but layouts/paddings aren't adapted. Systematic pass needed.",
     "/adapt"),
    ("P1-8", "PreParty page hides chat and channels on mobile",
     "preparty/page.tsx:122 and :307 (hidden md:flex)",
     "Responsive",
     "—",
     "Bottom-sheet or tab pattern for chat/channels on mobile.",
     "/adapt"),
    ("P1-9", "Liquid-glass pattern inlined in 15+ locations",
     "ApexScanOverlay.tsx, PreviewPanel.tsx:28,38,45,324,331, PillButton.tsx:18-25 (the one correct usage), + form pages",
     "Theming / Anti-Pattern",
     "—",
     "Promote to a single .liquid-glass CSS class in globals.css and/or a <LiquidGlassPill> component.",
     "/distill"),
    ("P1-10", "Gradient-text pattern duplicated 15×",
     "Search WebkitBackgroundClip:'text' — 15 hits across ApexScanOverlay, agents page, preparty, headers",
     "Theming",
     "—",
     "Extract <GradientText> component (accepts className for font choice) or .gradient-text-brand utility class.",
     "/distill"),
    ("P1-11", "PillButton under-used",
     "Defined in PillButton.tsx, used only in FilterPills.tsx:11. Inline liquid-glass pills exist in ApexScanOverlay.tsx:204, PreviewPanel.tsx:38,45, and several form pages",
     "Theming / Code Reuse",
     "—",
     "Extend PillButton API (variants: active / default / danger) and adopt across the 4+ inline-pill sites.",
     "/distill"),
]

for tag, title, location, category, wcag, recommendation, cmd in p1_findings:
    add_finding(doc, tag, P1, title, [
        ("Location", location),
        ("Category", category),
        ("WCAG/Standard", wcag),
        ("Recommendation", recommendation),
        ("Suggested command", cmd),
    ])

# ------------------------------------------------------------
# P2
# ------------------------------------------------------------
add_heading(doc, "P2 — Minor (address in next pass)", level=2, color=P2)

p2_bullets = [
    "[P2-1] SideNav.tsx uses <aside> but nav links aren't wrapped in <nav> landmark (a11y).",
    "[P2-2] Modals don't have focus trap or Escape-to-close in a documented/repeatable way — PreviewPanel.tsx mobile bottom-sheet lacks role=\"dialog\" and aria-modal.",
    "[P2-3] Index-based .map() keys on dynamic lists in ApexScanOverlay.tsx:216, 244, 307, 389, 446.",
    "[P2-4] Inline style mutations on hover in CommunityCard.tsx:17-18 instead of CSS classes.",
    "[P2-5] Touch targets <44px — w-8 h-8 and w-9 h-9 buttons exist in scan overlay, carousel arrows, modals (acceptable on desktop, borderline on mobile since these pages render on small screens).",
    "[P2-6] Timezone selector pill in slot-selection is a styled div, not a button — can't be activated by keyboard (slot-selection/page.tsx:357-364).",
    "[P2-7] Avatar alt=\"Profile\" and alt=\"Ajay Kumar\" — generic/redundant alt text in TopNav.tsx:55, OutreachDraftModal.tsx:184.",
    "[P2-8] Nine distinct 'orange' hex values for what should be ~3 tokens (see P0-3 detail).",
]
for b in p2_bullets:
    add_bullet(doc, b)

# ------------------------------------------------------------
# P3
# ------------------------------------------------------------
add_heading(doc, "P3 — Polish", level=2, color=P3)

p3_bullets = [
    "[P3-1] Decorative elements (brand swatches, dots) exposed to screen readers — should have aria-hidden=\"true\".",
    "[P3-2] Minor duplicate imports; a few unused.",
    "[P3-3] animate-apex-pulse uses transform: scale() (correct), but .glass-button:hover uses translateY(-1px) — fine, noting.",
    "[P3-4] No skip-to-main-content link in the layout.",
    "[P3-5] No lang override on any content; English-only is fine for now.",
]
for b in p3_bullets:
    add_bullet(doc, b)

# ------------------------------------------------------------
# PATTERNS & SYSTEMIC ISSUES
# ------------------------------------------------------------
add_heading(doc, "Patterns & Systemic Issues", level=1, color=PRIMARY)

doc.add_paragraph(
    "These three account for most of the findings and represent systemic gaps rather than one-offs:"
)

systemic = [
    ('"Documented but not extracted"',
     "The CLAUDE.md lists liquid-glass, gradient-text, and other patterns with exact inline styles to copy. This is the inverse of how a design system should work. Every screen built this way compounds the theming debt. Highest leverage fix."),
    ('"Icon buttons assume sighted, mouse-driven users"',
     "The app has no discipline around aria-label on icon-only controls. This will keep regressing with every new component unless an <IconButton> primitive is introduced that makes aria-label required at the prop-types level."),
    ('"Sign-up funnel treated as screenshot, not screen"',
     "The BARE_ROUTES pages exist as mirror-of-design implementations with fixed pixel widths. They bypass the responsive AppShell. For the demo this is fine; for production users to actually sign up, they'll need a systematic mobile pass."),
]
for i, (heading, body) in enumerate(systemic, 1):
    p = doc.add_paragraph()
    p.add_run(f"{i}. ").bold = True
    hr = p.add_run(heading + " — ")
    hr.bold = True
    hr.font.color.rgb = ACCENT
    p.add_run(body)

# ------------------------------------------------------------
# POSITIVE FINDINGS
# ------------------------------------------------------------
add_heading(doc, "Positive Findings (keep doing these)", level=1, color=PRIMARY)

positives = [
    ("Test-per-component discipline", "every component in components/ has a sibling __tests__/ file. This is rare at this stage and valuable."),
    ("RAF and timer cleanup is correct", "ApexScanOverlay.tsx and SuccessToast.tsx properly cancel/clear on unmount. No leaks."),
    ("Shared primitives exist", "PillButton, PrimaryButton, OutlineButton, SuccessToast are well-designed. The problem is under-adoption, not the primitives."),
    ("Tailwind token palette is well-structured", "Material You-style token names (primary, on-background, surface-container-*) are thoughtful. Issue is inconsistent use, not design."),
    ("Modern stack", "Next.js 16, React 19, TypeScript strict, Tailwind v3. No legacy baggage."),
    ("AppShell architecture is sound", "BARE_ROUTES, useIsMobile, <BottomNav>/<SideNav> split, the concave-corner detail all reflect real design work."),
    ("CLAUDE.md is a genuine asset", "the gotchas section (tooltip portals, padding-bottom in scroll containers, key-based overlay resets) captures hard-won knowledge future contributors would otherwise rediscover. Keep investing here."),
    ("No API coupling for demo data", "fixtures in lib/fixtures/ are cleanly typed and isolated. When real APIs arrive, this layer is the integration point."),
]
for heading, body in positives:
    p = doc.add_paragraph(style='List Bullet')
    hr = p.add_run(heading + " — ")
    hr.bold = True
    p.add_run(body)

# ------------------------------------------------------------
# RECOMMENDED ACTIONS
# ------------------------------------------------------------
add_heading(doc, "Recommended Actions (priority order)", level=1, color=PRIMARY)

actions = [
    ("P0", "/harden", "Fix OutreachDraftModal dialog semantics, add aria-labels to 11 icon buttons, restore focus indicators on 4 forms, link form labels to inputs (P0-1, P1-1, P1-2, P1-4)"),
    ("P0", "/optimize", "Replace CDN fonts with next/font/local, migrate 13 <img> tags to next/image, debounce the carousel resize listener (P0-2, P1-5, P1-6)"),
    ("P1", "/distill", "Extract liquid-glass, gradient-text, and scan-pill into shared primitives/classes; extend PillButton to cover the 4+ inline-pill sites (P1-9, P1-10, P1-11) — this is the single highest-ROI systemic fix"),
    ("P1", "/adapt", "Responsive pass on BARE_ROUTES signup funnel and on the PreParty mobile layout (P1-7, P1-8)"),
    ("P1", "/clarify", "Route or disable the four href=\"#\" dead links (P1-3)"),
    ("P2", "/harden", "Second pass: <nav> landmarks in SideNav/BottomNav, index-key cleanup, focus-trap utility for modals"),
    ("P2", "/polish", "Token migration: consolidate the 9 orange hex variants to named tokens, reduce 148 inline hex colors"),
]
for i, (sev, cmd, body) in enumerate(actions, 1):
    p = doc.add_paragraph(style='List Number')
    sev_run = p.add_run(f"[{sev}] ")
    sev_run.bold = True
    sev_run.font.color.rgb = P0 if sev == "P0" else (P1 if sev == "P1" else P2)
    cmd_run = p.add_run(cmd + " — ")
    cmd_run.bold = True
    cmd_run.font.name = "Consolas"
    p.add_run(body)

# ------------------------------------------------------------
# BOTTOM LINE
# ------------------------------------------------------------
add_heading(doc, "Bottom Line", level=1, color=PRIMARY)

doc.add_paragraph(
    "The codebase is a 5/10 as a production foundation. The architecture (Next.js 16, tests, token "
    "system, shared primitives, documented patterns, AppShell) is genuinely good — that's the 5 points "
    "you have. What's missing (a11y discipline, image optimization, mobile parity for sign-up, "
    "extraction of documented patterns) is the 5 points you don't. Critically, none of the gaps are "
    "structural — they're all mechanical hardening work. Run /harden + /distill + /adapt before adding "
    "more screens, because every new screen currently multiplies these gaps linearly."
)

# Save
doc.save(str(OUT))
print(f"Wrote: {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")
